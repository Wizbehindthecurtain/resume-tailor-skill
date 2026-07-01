#!/usr/bin/env python3
"""Render resume Markdown to a polished, single-column ATS-friendly DOCX, and extract text back.

Styling is "navy modern": a deep-navy name and uppercased section headers with a hairline
rule, on a clean Calibri body. All styling is ATS-safe — single column, real text, no
tables/images/text-boxes/columns; color and borders are font/paragraph formatting that
parsers ignore while still reading the text.
"""
import argparse
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- design tokens (navy modern) ---
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)
BODY_FONT = "Calibri"
NAVY_HEX = "1F3A5F"


def _bottom_rule(paragraph, color_hex=NAVY_HEX):
    """Add a hairline bottom border to a paragraph (ATS-safe — paragraph formatting)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _style_run(run, size, *, bold=False, color=None):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _spacing(paragraph, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def render_resume(markdown_text: str, out_path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)

    just_after_name = False
    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if not line:
            just_after_name = False
            continue
        if line.startswith("### "):  # job / entry heading
            p = doc.add_paragraph()
            _style_run(p.add_run(line[4:]), 11, bold=True)
            _spacing(p, before=8, after=1)
            just_after_name = False
        elif line.startswith("## "):  # section header — uppercase navy + rule
            p = doc.add_paragraph()
            _style_run(p.add_run(line[3:].upper()), 11.5, bold=True, color=NAVY)
            _spacing(p, before=12, after=3)
            _bottom_rule(p)
            just_after_name = False
        elif line.startswith("# "):  # name
            p = doc.add_paragraph()
            _style_run(p.add_run(line[2:]), 22, bold=True, color=NAVY)
            _spacing(p, before=0, after=1)
            just_after_name = True
        elif line.startswith("- "):  # bullet
            p = doc.add_paragraph(style="List Bullet")
            _style_run(p.add_run(line[2:]), 10.5)
            _spacing(p, before=0, after=1)
            just_after_name = False
        else:  # plain paragraph — contact line right after name, else body
            p = doc.add_paragraph()
            if just_after_name:
                _style_run(p.add_run(line), 9.5, color=GRAY)
                _spacing(p, before=0, after=6)
            else:
                _style_run(p.add_run(line), 10.5)
                _spacing(p, before=0, after=4)
            just_after_name = False

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def extract_text(docx_path) -> str:
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description="Render resume Markdown to DOCX, or extract text.")
    parser.add_argument("--in", dest="in_path", type=Path)
    parser.add_argument("--out", dest="out_path", type=Path)
    parser.add_argument("--extract", dest="extract_path", type=Path)
    args = parser.parse_args(argv)
    if args.extract_path:
        sys.stdout.write(extract_text(args.extract_path))
        return 0
    if args.in_path and args.out_path:
        render_resume(args.in_path.read_text(encoding="utf-8"), args.out_path)
        sys.stdout.write(str(args.out_path))
        return 0
    parser.error("provide --in and --out to render, or --extract to extract")


if __name__ == "__main__":
    raise SystemExit(main())
