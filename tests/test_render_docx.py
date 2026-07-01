import render_docx

def test_render_and_parse_back_keeps_keywords(tmp_path):
    md = ("# Jane Doe\n## Skills\n- Python (PYTHON)\n- Kubernetes\n"
          "## Work Experience\n### Engineer — Acme — 03/2022\n- Cut risk 40%\n")
    out = tmp_path / "r.docx"
    render_docx.render_resume(md, out)
    assert out.exists()
    text = render_docx.extract_text(out)
    for kw in ("Python", "Kubernetes", "Acme", "40%"):
        assert kw in text

def test_section_headers_styled_navy(tmp_path):
    from docx import Document
    out = tmp_path / "r.docx"
    render_docx.render_resume("# Jane Doe\n## Skills\n- Python\n", out)
    doc = Document(str(out))
    # find the SKILLS heading paragraph; its run should be navy (1F3A5F) and uppercased
    texts = [p.text for p in doc.paragraphs]
    assert "SKILLS" in texts  # section header uppercased
    para = next(p for p in doc.paragraphs if p.text == "SKILLS")
    run = para.runs[0]
    assert run.font.color and run.font.color.rgb is not None and str(run.font.color.rgb) == "1F3A5F"

def test_main_extract(tmp_path, capsys):
    out = tmp_path / "r.docx"
    render_docx.render_resume("# Jane\n- Python\n", out)
    rc = render_docx.main(["--extract", str(out)])
    assert rc == 0 and "Python" in capsys.readouterr().out
